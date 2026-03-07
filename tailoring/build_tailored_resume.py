import argparse
import re
import sqlite3
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

from tailoring.template_mapper import get_template_path, SECTION_ALIASES
from tailoring.job_keyword_ranker import extract_keywords, score_text_against_keywords

ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "database" / "jobs.db"
OUTPUT_DIR = ROOT / "output" / "resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# -----------------------------
# DB / job helpers
# -----------------------------

def slugify(text: str) -> str:
    text = (text or "").strip().lower()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")[:60] or "job"


def get_job(job_id: int) -> dict:
    con = sqlite3.connect(DB_PATH)
    row = con.execute(
        """
        SELECT id, title, company, description, track, url, apply_url
        FROM jobs
        WHERE id = ?
        """,
        (job_id,),
    ).fetchone()
    con.close()

    if not row:
        raise ValueError(f"Job id {job_id} not found")

    return {
        "id": row[0],
        "title": row[1] or "",
        "company": row[2] or "",
        "description": row[3] or "",
        "track": row[4] or "unknown",
        "url": row[5] or "",
        "apply_url": row[6] or "",
    }


# -----------------------------
# Section helpers
# -----------------------------

def normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip()).upper()


def find_section_indices(doc: Document) -> dict:
    indices = {}
    alias_lookup = {
        logical: [a.upper() for a in aliases]
        for logical, aliases in SECTION_ALIASES.items()
    }

    for i, p in enumerate(doc.paragraphs):
        text = normalize_heading(p.text)
        for logical_name, aliases in alias_lookup.items():
            if text in aliases:
                indices[logical_name] = i
    return indices


def get_section_range(doc: Document, section_name: str) -> Tuple[int, int]:
    indices = find_section_indices(doc)
    if section_name not in indices:
        raise ValueError(f"Section '{section_name}' not found in template")

    start = indices[section_name]
    starts_sorted = sorted(indices.values())
    pos = starts_sorted.index(start)
    end = len(doc.paragraphs)
    if pos + 1 < len(starts_sorted):
        end = starts_sorted[pos + 1]
    return start, end


# -----------------------------
# Paragraph classification
# -----------------------------

BULLET_CHARS = ("•", "-", "–", "—", "*")


def is_bullet_text(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    return t.startswith(BULLET_CHARS)


def strip_bullet_prefix(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return t
    if t.startswith(BULLET_CHARS):
        return t[1:].strip()
    return t


def is_probable_role_header(text: str) -> bool:
    """
    Heuristic for lines like:
    - IT Support Specialist | Company | Dates
    - Software Engineer, Company
    """
    t = (text or "").strip()
    if not t:
        return False
    if is_bullet_text(t):
        return False

    # role headers often include separators or date-ish text
    if "|" in t:
        return True
    if re.search(r"\b(20\d{2}|Present|Remote|CA|TX|Intern|Engineer|Analyst|Developer|Specialist|Research|Assistant)\b", t, re.I):
        return True
    return False


# -----------------------------
# Summary handling
# -----------------------------

TRACK_SUMMARY_MAP = {
    "data": (
        "Data-focused professional with experience in Python, SQL, analytics workflows, "
        "dashboarding, and cloud-based data solutions. Background includes building reliable "
        "technical workflows, translating business needs into data outputs, and supporting "
        "decision-making with structured analysis."
    ),
    "it": (
        "IT and technical support professional with experience troubleshooting systems, "
        "supporting users, documenting workflows, and working with cloud and software tools. "
        "Brings a strong mix of technical problem-solving, communication, and process improvement."
    ),
    "software": (
        "Software-focused professional with experience in Python, application development, APIs, "
        "automation, and building technical solutions across academic and professional projects. "
        "Brings strong problem-solving skills and a practical approach to delivering working systems."
    ),
    "ml_ai": (
        "Machine learning and AI-focused professional with experience in Python, model development, "
        "data analysis, experimentation, and technical research. Background includes building practical "
        "solutions and applying analytical methods to real-world problems."
    ),
    "unknown": (
        "Technical professional with experience across software, data, and IT-oriented work, "
        "including problem solving, automation, and building practical solutions in structured environments."
    ),
}


def replace_summary(doc: Document, track: str):
    start, end = get_section_range(doc, "summary")
    body_indices = [i for i in range(start + 1, end) if doc.paragraphs[i].text.strip()]

    if not body_indices:
        return

    summary_text = TRACK_SUMMARY_MAP.get(track.lower(), TRACK_SUMMARY_MAP["unknown"])

    # Put summary into first body paragraph, clear the rest
    first_idx = body_indices[0]
    doc.paragraphs[first_idx].text = summary_text

    for idx in body_indices[1:]:
        doc.paragraphs[idx].text = ""


# -----------------------------
# Skills handling
# -----------------------------

def reorder_skill_line(line: str, keywords: list[str]) -> str:
    if ":" not in line:
        return line

    label, values = line.split(":", 1)
    items = [x.strip() for x in values.split(",") if x.strip()]
    if not items:
        return line

    def item_score(item: str) -> float:
        return score_text_against_keywords(item, keywords)

    ranked = sorted(items, key=item_score, reverse=True)
    return f"{label.strip()}: " + ", ".join(ranked)


def replace_skills(doc: Document, keywords: list[str]):
    start, end = get_section_range(doc, "skills")

    for i in range(start + 1, end):
        paragraph = doc.paragraphs[i]
        txt = paragraph.text.strip()

        if not txt or ":" not in txt:
            continue

        label, values = txt.split(":", 1)

        # reorder the skills like before
        reordered = reorder_skill_line(txt, keywords)
        label, values = reordered.split(":", 1)

        # clear paragraph
        paragraph.clear()

        # bold label
        run_label = paragraph.add_run(label.strip() + ":")
        run_label.bold = True

        # normal text for skills
        paragraph.add_run(" " + values.strip())


# -----------------------------
# Experience handling
# -----------------------------

def score_bullet(bullet_text: str, keywords: list[str], role_header: str = "") -> float:
    score = score_text_against_keywords(bullet_text, keywords)
    score += 0.35 * score_text_against_keywords(role_header, keywords)
    return score


def trim_experience_section(doc: Document, keywords: list[str], track: str):
    start, end = get_section_range(doc, "experience")

    # Collect role blocks:
    # role header + contiguous bullets until next header/blank/section end
    i = start + 1
    role_blocks = []

    while i < end:
        current_text = doc.paragraphs[i].text.strip()

        if not current_text:
            i += 1
            continue

        if is_probable_role_header(current_text):
            role_idx = i
            role_header = current_text
            bullet_indices = []

            j = i + 1
            while j < end:
                t = doc.paragraphs[j].text.strip()
                if not t:
                    j += 1
                    continue
                if is_probable_role_header(t) and not is_bullet_text(t):
                    break
                if is_bullet_text(t):
                    bullet_indices.append(j)
                j += 1

            role_blocks.append((role_idx, role_header, bullet_indices))
            i = j
        else:
            i += 1

    # If we can't identify role blocks, do nothing
    if not role_blocks:
        return

    # One-page-oriented bullet budget
    # newest/first role gets more bullets
    role_limits = []
    for idx, _ in enumerate(role_blocks):
        if idx == 0:
            role_limits.append(4 if track.lower() in {"it", "software"} else 5)
        elif idx == 1:
            role_limits.append(2)
        else:
            role_limits.append(1)

    # Trim bullets inside each role
    for block_idx, (role_idx, role_header, bullet_indices) in enumerate(role_blocks):
        if not bullet_indices:
            continue

        limit = role_limits[min(block_idx, len(role_limits) - 1)]

        scored = []
        for b_idx in bullet_indices:
            raw_text = doc.paragraphs[b_idx].text
            bullet_text = strip_bullet_prefix(raw_text)
            s = score_bullet(bullet_text, keywords, role_header)
            scored.append((s, b_idx, raw_text))

        scored_sorted = sorted(scored, key=lambda x: x[0], reverse=True)
        keep = {b_idx for _, b_idx, _ in scored_sorted[:limit]}

        # Clear unwanted bullets, preserve wanted ones exactly as written
        for _, b_idx, _ in scored:
            if b_idx not in keep:
                doc.paragraphs[b_idx].text = ""

    # Collapse obvious multi-blank runs within experience by clearing extra blanks
    blank_run = 0
    for i in range(start + 1, end):
        if not doc.paragraphs[i].text.strip():
            blank_run += 1
            if blank_run > 1:
                doc.paragraphs[i].text = ""
        else:
            blank_run = 0


# -----------------------------
# Save
# -----------------------------

def save_doc(doc: Document, job: dict) -> Path:
    company_slug = slugify(job["company"])
    title_slug = slugify(job["title"])
    track_slug = slugify(job["track"])
    out = OUTPUT_DIR / f"{job['id']}_{company_slug}_{title_slug}_{track_slug}.docx"
    doc.save(out)
    return out


# -----------------------------
# Main
# -----------------------------

def main(job_id: int):
    job = get_job(job_id)
    template_path = get_template_path(job["track"])
    doc = Document(template_path)

    jd_text = f"{job['title']} {job['description']}"
    keywords = extract_keywords(jd_text, top_n=25)

    replace_summary(doc, job["track"])
    replace_skills(doc, keywords)
    trim_experience_section(doc, keywords, job["track"])

    out_path = save_doc(doc, job)
    print(f"Created tailored resume: {out_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--job_id", type=int, required=True)
    args = parser.parse_args()
    main(args.job_id)